import streamlit as st
import requests
import time
from io import BytesIO
import os


# Set page config to collapse sidebar by default
st.set_page_config(initial_sidebar_state="collapsed")

# --- Configuration ---
FASTAPI_BASE_URL = os.getenv("FASTAPI_BASE_URL", "http://localhost:8001/api/v1")
TRANSCRIPTION_ENDPOINT = f"{FASTAPI_BASE_URL}/transcribe"
SUMMARIZATION_ENDPOINT = f"{FASTAPI_BASE_URL}/llm/summarize"
ACTION_ITEMS_ENDPOINT = f"{FASTAPI_BASE_URL}/llm/extract-action-items"
CHAT_ENDPOINT = f"{FASTAPI_BASE_URL}/llm/chat"

# --- Page Setup ---
st.set_page_config(
    page_title="PolyNote",
    page_icon="📝",
    layout="wide",
)

# --- Styling (Optional) ---
st.markdown("""
<style>
    /* Modern dark theme with new color scheme */
    .stApp {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        color: #ffffff;
        scroll-behavior: smooth;
        font-family: 'Inter', sans-serif;
    }
    
    /* Custom scrollbar with new colors */
    ::-webkit-scrollbar {
        width: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: rgba(255, 255, 255, 0.05);
    }
    
    ::-webkit-scrollbar-thumb {
        background: linear-gradient(45deg, #3b82f6, #8b5cf6);
        border-radius: 4px;
        box-shadow: 0 0 10px rgba(59, 130, 246, 0.5);
    }
    
    /* Header styling with glassmorphism */
    .stApp > header {
        background: rgba(30, 41, 59, 0.8);
        backdrop-filter: blur(10px);
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: 0 4px 30px rgba(0, 0, 0, 0.1);
    }
    
    /* Animated title with enhanced blur effect */
    h1 {
        text-align: center;
        color: #8b5cf6;
        font-weight: 800;
        position: relative;
        animation: titleFloat 3s ease-in-out infinite;
        z-index: 1;
        text-shadow: 0 0 30px rgba(139, 92, 246, 0.6), 
                     0 0 60px rgba(59, 130, 246, 0.4);
    }
    
    h1::before {
        content: '🎙️ PolyNote';
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        filter: blur(12px);
        opacity: 0.6;
        z-index: -1;
        animation: titleBlur 3s ease-in-out infinite;
        background: linear-gradient(45deg, #3b82f6, #8b5cf6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
    }
    
    @keyframes titleFloat {
        0%, 100% { transform: translateY(0); }
        50% { transform: translateY(-10px); }
    }
    
    @keyframes titleBlur {
        0%, 100% { filter: blur(12px); opacity: 0.6; }
        50% { filter: blur(16px); opacity: 0.4; }
    }
    
    /* Modern button styling with new gradient */
    .stButton>button {
        border-radius: 12px;
        border: none;
        color: #ffffff;
        background: linear-gradient(45deg, #3b82f6, #8b5cf6);
        transition: all 0.3s ease-in-out;
        margin-bottom: 10px;
        box-shadow: 0 4px 15px rgba(59, 130, 246, 0.2);
        position: relative;
        overflow: hidden;
    }
    
    .stButton>button:before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(
            90deg,
            transparent,
            rgba(255, 255, 255, 0.2),
            transparent
        );
        transition: 0.5s;
    }
    
    .stButton>button:hover {
        transform: translateY(-3px);
        box-shadow: 0 6px 20px rgba(59, 130, 246, 0.4);
    }
    
    .stButton>button:hover:before {
        left: 100%;
    }
    
    /* File uploader with new styling */
    .stFileUploader {
        border: 2px dashed #3b82f6;
        border-radius: 12px;
        padding: 2rem;
        background: rgba(59, 130, 246, 0.05);
        transition: all 0.3s ease-in-out;
        position: relative;
        overflow: hidden;
    }
    
    .stFileUploader:before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        width: 100%;
        height: 100%;
        background: linear-gradient(45deg, transparent, rgba(59, 130, 246, 0.1), transparent);
        transform: translateX(100%);
        transition: 0.5s;
    }
    
    .stFileUploader:hover:before {
        transform: translateX(-100%);
    }
    
    /* Chat message styling */
    .stChatMessage {
        border-radius: 12px;
        padding: 1rem;
        margin-bottom: 1rem;
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.1);
        transition: all 0.3s ease-in-out;
        animation: slideIn 0.5s ease-out;
    }
    
    @keyframes slideIn {
        from {
            transform: translateX(-20px);
            opacity: 0;
        }
        to {
            transform: translateX(0);
            opacity: 1;
        }
    }
    
    /* Input field styling */
    .stTextInput>div>div>input,
    .stTextArea>div>div>textarea {
        background: rgba(255, 255, 255, 0.05);
        color: #ffffff;
        border: 1px solid #3b82f6;
        border-radius: 12px;
        padding: 0.8rem 1rem;
        transition: all 0.3s ease-in-out;
    }
    
    .stTextInput>div>div>input:focus,
    .stTextArea>div>div>textarea:focus {
        background: rgba(255, 255, 255, 0.1);
        box-shadow: 0 0 15px rgba(59, 130, 246, 0.3);
        border-color: #8b5cf6;
    }
    
    /* Sidebar styling */
    .css-1d391kg {
        background: rgba(30, 41, 59, 0.8);
        backdrop-filter: blur(10px);
        border-right: 1px solid rgba(255, 255, 255, 0.1);
        box-shadow: 4px 0 30px rgba(0, 0, 0, 0.2);
    }
    
    /* Loading spinner animation */
    .stSpinner > div {
        background: linear-gradient(45deg, #3b82f6, #8b5cf6);
    }
    
    /* Divider with new gradient */
    hr {
        border: none;
        height: 1px;
        background: linear-gradient(90deg, transparent, #3b82f6, #8b5cf6, transparent);
        margin: 2rem 0;
    }
    
    /* Card hover effect */
    .element-container {
        transition: all 0.3s ease-in-out;
    }
    
    .element-container:hover {
        transform: translateY(-5px);
    }
    
    /* Success/Error message styling */
    .stAlert {
        background: rgba(255, 255, 255, 0.05);
        backdrop-filter: blur(10px);
        border-radius: 12px;
        border: 1px solid rgba(255, 255, 255, 0.1);
        animation: fadeIn 0.5s ease-out;
    }
    
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(-10px); }
        to { opacity: 1; transform: translateY(0); }
    }
</style>
""", unsafe_allow_html=True)

# --- Session State Initialization ---
def init_session_state():
    defaults = {
        'transcript_data': None,
        'error_message': None,
        'is_loading': False,
        'uploaded_filename': None,
        'summary_data': None,
        'summarizing': False,
        'action_items_data': None,
        'summary_error': None,
        'chat_history': [],
        'chatting': False,
        'chat_error': None,
        'full_transcript_text': None
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val

init_session_state()

# --- Helper Functions ---
from fpdf import FPDF
import pandas as pd
from docx import Document

def reset_state():
    keys_to_skip = ['audio_uploader', 'chat_input', 'transcribe_button', 'summarize_btn']  # widget keys to skip
    for k in st.session_state.keys():
        if k in keys_to_skip:
            continue
        if k == 'chat_history':
            st.session_state[k] = []
        elif k in ['is_loading', 'summarizing', 'chatting']:
            st.session_state[k] = False
        else:
            st.session_state[k] = None

def highlight_text(text, query):
    if not query:
        return text
    import re
    pattern = re.compile(re.escape(query), re.IGNORECASE)
    # Highlight with HTML span and yellow background
    return pattern.sub(lambda m: f'<span style="background-color: #ffd700; color: #000000; padding: 2px 4px; border-radius: 3px; font-weight: bold;">{m.group(0)}</span>', text)

def generate_txt(summary_data: dict, transcript_text: str) -> BytesIO:
    buffer = BytesIO()
    summary = summary_data.get("summary", "No summary available")
    action_items = summary_data.get("action_items", [])
    action_items_text = "\n".join([f"- {item}" for item in action_items]) if action_items else "No specific action items identified."
    content = f"""Meeting Notes\n\nSummary:\n{summary}\n\nAction Items:\n{action_items_text}\n\nTranscript:\n{transcript_text or 'No transcript available.'}"""
    buffer.write(content.strip().encode("utf-8"))
    buffer.seek(0)
    return buffer

def generate_pdf(summary_data: dict, transcript_text: str) -> BytesIO:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Meeting Notes", ln=True, align='C')
    pdf.ln(10)
    pdf.multi_cell(0, 10, txt=f"Summary:\n{summary_data.get('summary', 'No summary available')}")
    pdf.ln(5)
    action_items = summary_data.get("action_items", [])
    action_items_text = "\n".join([f"- {item}" for item in action_items]) if action_items else "No specific action items identified."
    pdf.multi_cell(0, 10, txt=f"Action Items:\n{action_items_text}")
    pdf.ln(5)
    pdf.multi_cell(0, 10, txt=f"Transcript:\n{transcript_text or 'No transcript available.'}")
    # fpdf output() returns bytes when dest='S'
    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    buffer = BytesIO(pdf_bytes)
    buffer.seek(0)
    return buffer

def generate_csv(summary_data: dict, transcript_text: str) -> BytesIO:
    df = pd.DataFrame({
        "Summary": [summary_data.get('summary', 'No summary available')],
        "Action Items": ["; ".join(summary_data.get('action_items', []))],
        "Transcript": [transcript_text or 'No transcript available.']
    })
    buffer = BytesIO()
    df.to_csv(buffer, index=False)
    buffer.seek(0)
    return buffer

def generate_docx(summary_data: dict, transcript_text: str) -> BytesIO:
    doc = Document()
    doc.add_heading("Meeting Notes", 0)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary_data.get('summary', 'No summary available'))
    doc.add_heading("Action Items", level=1)
    action_items = summary_data.get("action_items", [])
    if action_items:
        for item in action_items:
            doc.add_paragraph(item, style='List Bullet')
    else:
        doc.add_paragraph("No specific action items identified.")
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(transcript_text or 'No transcript available.')
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- UI Layout ---
st.title("PolyNote")
st.markdown("---")

with st.sidebar:
    st.header("Upload Audio File")

    uploaded_file = st.file_uploader(
        "Choose an audio file...",
        type=['mp3', 'wav', 'm4a', 'ogg', 'mp4', 'webm', 'mov', 'flac', 'mkv'],
        key="audio_uploader",
        on_change=reset_state,
        help="Upload your recording"
    )

    if uploaded_file:
        # Store filename for display
        if st.session_state.uploaded_filename != uploaded_file.name:
            st.session_state.uploaded_filename = uploaded_file.name
        st.info(f"File selected: `{st.session_state.uploaded_filename}`")

        # Transcription Button
        transcribe_disabled = st.session_state.is_loading or st.session_state.transcript_data is not None
        if st.button("Transcribe Audio", key="transcribe_button", use_container_width=True, disabled=transcribe_disabled):
            reset_state()
            st.session_state.transcribe_clicked = True
            st.session_state.uploaded_filename = uploaded_file.name
            st.session_state.is_loading = True
            st.write("Transcribing audio... This might take a moment depending on the file size.")
            try:
                files = {'file': (uploaded_file.name, uploaded_file.getvalue(), uploaded_file.type)}
                resp = requests.post(TRANSCRIPTION_ENDPOINT, files=files)
                resp.raise_for_status()
                result = resp.json()
                status = result.get("status")
                if status == "TranscriptStatus.completed":
                    st.session_state.transcript_data = result
                    st.session_state.full_transcript_text = result.get('text', '')
                else:
                    err = result.get('error', 'Unknown error')
                    st.session_state.error_message = f"Transcription failed: {err}"
            except Exception as e:
                st.session_state.error_message = f"API Error: {e}"
            finally:
                st.session_state.is_loading = False
                st.rerun()

        # Summarization Button
        can_summarize = bool(st.session_state.full_transcript_text)
        sum_disabled = not can_summarize or st.session_state.summarizing
        if st.button("Summarize Transcript", key="summarize_btn", use_container_width=True, disabled=sum_disabled):
            st.session_state.summarizing = True
            st.write("Generating summary and extracting action items...")
            summary_data = None
            action_items_data = None
            action_items_error = None
            try:
                payload = {"transcript": st.session_state.full_transcript_text}
                # Call summary endpoint
                summary_resp = requests.post(SUMMARIZATION_ENDPOINT, json=payload, timeout=180)
                summary_resp.raise_for_status()
                summary_data = summary_resp.json()
            except Exception as e:
                st.session_state.summary_error = f"Summarization Error: {e}"
            try:
                # Call action items endpoint
                action_resp = requests.post(ACTION_ITEMS_ENDPOINT, json=payload, timeout=180)
                action_resp.raise_for_status()
                action_items_data = action_resp.json()
            except Exception as e:
                action_items_error = f"Action Items Error: {e}"
            # Merge action items if available
            if summary_data:
                summary_data["action_items"] = action_items_data.get("action_items", []) if action_items_data else []
                st.session_state.summary_data = summary_data
                st.session_state.summary_error = None
                st.session_state.action_items_error = action_items_error
            elif action_items_error:
                st.session_state.summary_error = action_items_error
            st.session_state.summarizing = False
            st.rerun()

    # Status Display
    st.markdown("---")
    st.subheader("Search in Transcript")
    search_query = st.text_input("Enter a keyword to search", placeholder="e.g., budget, next meeting, action", key="keyword_search")
    
    st.subheader("Status")
    if st.session_state.is_loading:
        st.info("Transcription in progress...")
    elif st.session_state.summarizing:
        st.info("Summarization in progress...")
    elif st.session_state.chatting:
        st.info("Waiting for chat response...")
    elif st.session_state.error_message:
        st.error(st.session_state.error_message)
    elif st.session_state.summary_error:
        st.error(st.session_state.summary_error)
    elif st.session_state.transcript_data:
        st.success("Ready")
    else:
        st.info("Upload a file to start.")


if st.session_state.get("transcribe_clicked"):
    data = st.session_state.transcript_data
    st.subheader("Meeting Transcript")
    # --- Display Detected Language ---
    if data and data.get('language_code'):
        st.info(f"Detected Language: {data['language_code']}")
    # Display transcript with keyword filter
    utterances = data.get('utterances', []) if data else []
    if utterances:
        # Filter utterances if search query exists
        filtered_utterances = utterances
        if search_query:
            filtered_utterances = [utt for utt in utterances if search_query.lower() in utt.get('text', '').lower()]
            if not filtered_utterances:
                st.warning(f"No matches found for '{search_query}'")
            else:
                st.info(f"Found {len(filtered_utterances)} matching utterances")
        
        for utt in filtered_utterances:
            speaker = utt.get('speaker', 'Unknown')
            start_ms = utt.get('start', 0)
            timestamp = time.strftime('%M:%S', time.gmtime(start_ms/1000))
            text = utt.get('text', '')
            with st.chat_message(speaker):
                if search_query:
                    st.markdown(f"_{timestamp}_ | {highlight_text(text, search_query)}", unsafe_allow_html=True)
                else:
                    st.markdown(f"_{timestamp}_ | {text}")
    else:
        # fallback for full transcript
        if search_query:
            matched_lines = [line for line in st.session_state.full_transcript_text.split('\n') if search_query.lower() in line.lower()]
            if matched_lines:
                st.write("### Matching Lines:")
                for line in matched_lines:
                    st.markdown(highlight_text(line, search_query), unsafe_allow_html=True)
            else:
                st.warning("No matches found.")
        else:
            st.text_area("Full Transcript", value=st.session_state.full_transcript_text, height=300, disabled=True)

    st.markdown("---")

    if st.session_state.summary_data:
        sd = st.session_state.summary_data
        st.subheader("Meeting Summary")
        st.markdown(sd.get('summary','_No summary generated._'))
        st.subheader("Action Items")
        items = sd.get('action_items', [])
        if items:
            for it in items:
                st.markdown(f"- {it}")
        else:
            st.info("No specific action items identified.")
        if st.session_state.get("action_items_error"):
            st.warning(st.session_state.action_items_error)
        txt_buffer = generate_txt(sd, st.session_state.full_transcript_text)
        st.download_button(
            label="Export to txt",
            data=txt_buffer,
            file_name=f"{st.session_state.uploaded_filename.split('.')[0]}.txt",
            mime="application/txt",
            use_container_width=True
        )
        # --- Export PDF, CSV, DOCX ---
        pdf_buffer = generate_pdf(sd, st.session_state.full_transcript_text)
        st.download_button(
            label="Export to PDF",
            data=pdf_buffer,
            file_name=f"{st.session_state.uploaded_filename.split('.')[0]}.pdf",
            mime="application/pdf",
            use_container_width=True
        )
        csv_buffer = generate_csv(sd, st.session_state.full_transcript_text)
        st.download_button(
            label="Export to CSV",
            data=csv_buffer,
            file_name=f"{st.session_state.uploaded_filename.split('.')[0]}.csv",
            mime="text/csv",
            use_container_width=True
        )
        docx_buffer = generate_docx(sd, st.session_state.full_transcript_text)
        st.download_button(
            label="Export to DOCX",
            data=docx_buffer,
            file_name=f"{st.session_state.uploaded_filename.split('.')[0]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        st.markdown("---")

    # Chat Interface
    st.subheader("Chat with Transcript Context")
    for msg in st.session_state.chat_history:
        role = msg['role']
        with st.chat_message(role):
            st.markdown(msg['content'])

    user_input = st.chat_input("Ask a question about the transcript...", disabled=st.session_state.chatting)
    if user_input:
        st.session_state.chat_history.append({"role": "user", "content": user_input})
        st.session_state.chatting = True
        st.rerun()

    # On rerun, if chatting and last message is user, call Chat API
    if st.session_state.chatting and st.session_state.chat_history[-1]['role'] == 'user':
        with st.spinner("Thinking..."):
            try:
                payload = {
                    "transcript_context": st.session_state.full_transcript_text,
                    "user_query": st.session_state.chat_history[-1]['content']
                }
                resp = requests.post(CHAT_ENDPOINT, json=payload, timeout=120)
                resp.raise_for_status()
                ai_resp = resp.json().get('ai_response','')
                st.session_state.chat_history.append({"role": "assistant", "content": ai_resp})
            except Exception as e:
                st.session_state.chat_history.append({"role": "assistant", "content": f"Error: {e}"})
            finally:
                st.session_state.chatting = False
                st.rerun()
else:
    if not st.session_state.is_loading and not st.session_state.error_message:
        st.info("Upload an audio file and click 'Transcribe Audio' to begin.")

# --- Footer ---
st.markdown("---")
st.caption("Copyright 2025 Vipul Pawar")
